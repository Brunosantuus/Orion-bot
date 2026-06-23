import json
import os
import re
import schedule
import tempfile
import threading
import time
import sys
import telebot
import requests
from datetime import date, datetime, timedelta
from zoneinfo import ZoneInfo

FUSO = ZoneInfo("America/Sao_Paulo")

_base = os.path.dirname(__file__)
CONFIG_FILE = (os.path.join(_base, "config.json")
               if os.path.exists(os.path.join(_base, "config.json"))
               else os.path.join(_base, "config.railway.json"))
TAREFAS_FILE = os.path.join(_base, "tarefas.json")
LEMBRETES_USUARIO_FILE = os.path.join(_base, "lembretes_usuario.json")
STREAKS_FILE = os.path.join(_base, "streaks.json")
NOTAS_FILE = os.path.join(_base, "notas.json")
MEMORIA_FILE = os.path.join(_base, "memoria.json")
PESO_FILE = os.path.join(_base, "peso_historico.json")
AGUA_LOG_FILE = os.path.join(_base, "agua_log.json")

config = {}
bot = None
historico_chat = {}   # cache em memória: {chat_id: [mensagens]}
estados = {}
_snooze_pending = {}

HIST_MAX = 40         # mensagens salvas por conversa no Supabase
HIST_CONTEXTO = 20    # mensagens enviadas à IA por chamada

_HABITO_POR_HORARIO = {"10:00": "agua", "15:00": "agua", "18:00": "academia"}

# Saquarema RJ
_LAT = -22.9344
_LON = -42.5049

# Agenda padrão: 0=seg, 1=ter, 2=qua, 3=qui, 4=sex, 5=sab, 6=dom
_AGENDA_FIXA = {0: 'A', 1: 'B', 2: 'Descanso', 3: 'C', 4: 'A', 5: 'B', 6: 'Descanso'}

_TREINOS_FIXOS = {
    'A': [
        {'grupo': 'Peito',   'exercicio': 'Supino reto com barra',          'series': '4', 'repeticoes': '6-8'},
        {'grupo': 'Peito',   'exercicio': 'Supino inclinado com halteres',  'series': '3', 'repeticoes': '8-10'},
        {'grupo': 'Peito',   'exercicio': 'Crossover',                      'series': '3', 'repeticoes': '10-12'},
        {'grupo': 'Tríceps', 'exercicio': 'Paralelas',                      'series': '3', 'repeticoes': '6-10'},
        {'grupo': 'Tríceps', 'exercicio': 'Tríceps testa',                  'series': '3', 'repeticoes': '8-10'},
        {'grupo': 'Tríceps', 'exercicio': 'Tríceps corda no pulley',        'series': '3', 'repeticoes': '10-12'},
        {'grupo': 'Ombro',   'exercicio': 'Desenvolvimento de ombro',       'series': '3', 'repeticoes': '6-10'},
        {'grupo': 'Ombro',   'exercicio': 'Elevação lateral',               'series': '3', 'repeticoes': '10-15'},
    ],
    'B': [
        {'grupo': 'Quadríceps',    'exercicio': 'Agachamento livre',          'series': '4', 'repeticoes': '6-8'},
        {'grupo': 'Quadríceps',    'exercicio': 'Leg press',                  'series': '3', 'repeticoes': '8-12'},
        {'grupo': 'Quadríceps',    'exercicio': 'Cadeira extensora',          'series': '3', 'repeticoes': '10-12'},
        {'grupo': 'Posterior',     'exercicio': 'Mesa flexora',               'series': '3', 'repeticoes': '10-12'},
        {'grupo': 'Posterior',     'exercicio': 'Stiff',                      'series': '3', 'repeticoes': '8-10'},
        {'grupo': 'Panturrilha',   'exercicio': 'Panturrilha',                'series': '4', 'repeticoes': '12-15'},
        {'grupo': 'Adutor/Abdutor','exercicio': 'Cadeira abdutora/adutora',   'series': '3', 'repeticoes': '12-15'},
    ],
    'C': [
        {'grupo': 'Costas',  'exercicio': 'Barra fixa',                         'series': '4', 'repeticoes': '6-10'},
        {'grupo': 'Costas',  'exercicio': 'Puxada na frente',                   'series': '3', 'repeticoes': '8-12'},
        {'grupo': 'Costas',  'exercicio': 'Remada baixa',                       'series': '3', 'repeticoes': '8-12'},
        {'grupo': 'Costas',  'exercicio': 'Remada máquina com apoio no peito',  'series': '3', 'repeticoes': '10-12'},
        {'grupo': 'Costas',  'exercicio': 'Pullover',                           'series': '3', 'repeticoes': '10-12'},
        {'grupo': 'Bíceps',  'exercicio': 'Rosca direta',                       'series': '3', 'repeticoes': '8-10'},
        {'grupo': 'Bíceps',  'exercicio': 'Rosca alternada',                    'series': '3', 'repeticoes': '10-12'},
        {'grupo': 'Bíceps',  'exercicio': 'Rosca martelo',                      'series': '3', 'repeticoes': '10-12'},
        {'grupo': 'Abdômen', 'exercicio': 'Prancha',                            'series': '3', 'repeticoes': '30-60s'},
        {'grupo': 'Abdômen', 'exercicio': 'Elevação de pernas',                 'series': '3', 'repeticoes': '12-15'},
    ],
}

_NOME_DIA = {
    'A': 'Dia A — Peito + Tríceps + Ombro',
    'B': 'Dia B — Pernas + Panturrilha',
    'C': 'Dia C — Costas + Bíceps + Abdômen',
}

_ICONE_GRUPO = {
    'Peito': '🔵', 'Tríceps': '⚫', 'Ombro': '🟠',
    'Quadríceps': '🟢', 'Posterior': '🟤', 'Panturrilha': '🟡', 'Adutor/Abdutor': '🟣',
    'Costas': '🔵', 'Bíceps': '🟠', 'Abdômen': '⚫',
}


# ── Config ────────────────────────────────────────────────────────────────────

def carregar_config():
    global config
    with open(CONFIG_FILE, "r", encoding="utf-8") as f:
        config = json.load(f)
    token_id = os.environ.get("TOKEN_ID")
    token_key = os.environ.get("TOKEN_KEY")
    if token_id and token_key:
        config["telegram"]["token"] = f"{token_id}:{token_key}"
    if os.environ.get("TELEGRAM_CHAT_ID"):
        config["telegram"]["chat_id"] = os.environ["TELEGRAM_CHAT_ID"]
    if os.environ.get("GROQ_API_KEY"):
        config["groq"]["api_key"] = os.environ["GROQ_API_KEY"]


def salvar_config():
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=2)


# ── Supabase ──────────────────────────────────────────────────────────────────

_SB_URL = os.environ.get("SUPABASE_URL", "")
_SB_KEY = os.environ.get("SUPABASE_KEY", "")
_USA_SB = bool(_SB_URL and _SB_KEY)


def _sb_req(method, table, filters=None, body=None):
    qs = "&".join(f"{k}={v}" for k, v in (filters or {}).items())
    url = f"{_SB_URL}/rest/v1/{table}{'?' + qs if qs else ''}"
    headers = {
        "apikey": _SB_KEY,
        "Authorization": f"Bearer {_SB_KEY}",
        "Content-Type": "application/json",
        "Prefer": "return=representation",
    }
    r = requests.request(method, url, headers=headers, json=body, timeout=10)
    try:
        data = r.json() if r.content else []
        # Supabase devolve um dict em caso de erro (ex: tabela inexistente).
        # Os callers sempre esperam lista, então normalizamos.
        return data if isinstance(data, list) else []
    except Exception:
        return []


# ── Histórico de conversa (persistente) ──────────────────────────────────────

def _hist_carregar_db(chat_id):
    """Carrega histórico do Supabase para o cache em memória."""
    if not _USA_SB:
        return []
    try:
        dados = _sb_req("GET", "historico_chat", {
            "chat_id": f"eq.{chat_id}",
            "select": "role,content",
            "order": "id.asc",
            "limit": str(HIST_MAX),
        })
        return [{"role": d["role"], "content": d["content"]} for d in (dados or [])]
    except Exception:
        return []


def _hist_salvar(chat_id, role, content):
    """Salva mensagem no Supabase e descarta as mais antigas se passar do limite."""
    if not _USA_SB:
        return
    try:
        _sb_req("POST", "historico_chat", body={
            "chat_id": str(chat_id),
            "role": role,
            "content": content,
        })
        # Mantém só as últimas HIST_MAX mensagens
        todos = _sb_req("GET", "historico_chat", {
            "chat_id": f"eq.{chat_id}",
            "select": "id",
            "order": "id.asc",
        })
        if todos and len(todos) > HIST_MAX:
            for r in todos[:len(todos) - HIST_MAX]:
                _sb_req("DELETE", "historico_chat", {"id": f"eq.{r['id']}"})
    except Exception:
        pass


def _hist_limpar_db(chat_id):
    if not _USA_SB:
        return
    try:
        _sb_req("DELETE", "historico_chat", {"chat_id": f"eq.{chat_id}"})
    except Exception:
        pass


# ── Helpers de parsing ────────────────────────────────────────────────────────

def _parse_tarefa(texto):
    """Extrai prioridade (!urgente) e prazo (@DD/MM) do texto."""
    prioridade = "normal"
    prazo = None

    if texto.startswith("!") or re.search(r"\burgente\b", texto, re.IGNORECASE):
        prioridade = "alta"
        texto = re.sub(r"^!+\s*", "", texto)
        texto = re.sub(r"\burgente:?\s*", "", texto, flags=re.IGNORECASE).strip()

    match = re.search(r"@(\d{1,2})/(\d{1,2})(?:/(\d{4}))?", texto)
    if match:
        try:
            dia = int(match.group(1))
            mes = int(match.group(2))
            ano = int(match.group(3)) if match.group(3) else date.today().year
            prazo = date(ano, mes, dia).isoformat()
        except ValueError:
            pass
        texto = re.sub(r"@\d{1,2}/\d{1,2}(?:/\d{4})?", "", texto).strip()

    return texto, prioridade, prazo


def _prazo_label(prazo_iso):
    """Retorna label legível do prazo com urgência."""
    if not prazo_iso:
        return ""
    try:
        prazo_dt = date.fromisoformat(prazo_iso)
        dias = (prazo_dt - date.today()).days
        if dias < 0:
            return " ⚠️ VENCIDA"
        if dias == 0:
            return " 🚨 HOJE"
        if dias == 1:
            return " ⏳ amanhã"
        return f" 📅 {prazo_dt.strftime('%d/%m')}"
    except Exception:
        return ""


# ── Clima (Open-Meteo, sem API key) ──────────────────────────────────────────

def buscar_clima():
    """Retorna string com temperatura e condição em Saquarema ou None se falhar."""
    try:
        r = requests.get(
            "https://api.open-meteo.com/v1/forecast",
            params={
                "latitude": _LAT,
                "longitude": _LON,
                "current": "temperature_2m,weathercode",
                "hourly": "precipitation_probability",
                "forecast_days": 1,
                "timezone": "America/Sao_Paulo",
            },
            timeout=8,
        )
        if not r.ok:
            return None
        data = r.json()
        temp = data["current"]["temperature_2m"]
        wcode = data["current"]["weathercode"]

        _desc = {
            0: "céu limpo ☀️", 1: "quase limpo 🌤️", 2: "parcialmente nublado ⛅",
            3: "nublado ☁️", 45: "neblina 🌫️", 48: "neblina 🌫️",
            51: "garoa 🌦️", 53: "garoa 🌦️", 55: "garoa 🌦️",
            61: "chuva 🌧️", 63: "chuva 🌧️", 65: "chuva forte 🌧️",
            80: "pancadas 🌦️", 81: "pancadas 🌦️", 82: "pancadas fortes ⛈️",
            95: "tempestade ⛈️", 96: "tempestade ⛈️", 99: "tempestade ⛈️",
        }
        desc = _desc.get(wcode, "tempo variável")

        probs = data.get("hourly", {}).get("precipitation_probability", [])
        max_prob = max(probs[:12]) if probs else 0
        chuva = f", {max_prob}% de chuva" if max_prob >= 40 else ""

        return f"{temp:.0f}°C, {desc}{chuva}"
    except Exception:
        return None


# ── Peso e saúde ─────────────────────────────────────────────────────────────

def carregar_historico_peso():
    if _USA_SB:
        try:
            return _sb_req("GET", "peso_historico", {"select": "*", "order": "data.asc"}) or []
        except Exception:
            pass
    if os.path.exists(PESO_FILE):
        with open(PESO_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def _peso_local_save(lst):
    with open(PESO_FILE, "w", encoding="utf-8") as f:
        json.dump(lst, f, ensure_ascii=False, indent=2)


def registrar_peso(peso, observacao=None):
    novo = {"peso": peso, "data": date.today().isoformat(), "observacao": observacao}
    if _USA_SB:
        try:
            _sb_req("POST", "peso_historico", body=novo)
            return
        except Exception:
            pass
    lst = carregar_historico_peso()
    lst.append(novo)
    _peso_local_save(lst)


def peso_atual():
    historico = carregar_historico_peso()
    return float(historico[-1]["peso"]) if historico else None


def calcular_meta_agua(peso_kg):
    """35ml por kg de peso corporal."""
    return round(peso_kg * 0.035, 1)


def tendencia_peso(historico, n=5):
    """Retorna string com tendência de peso nos últimos n registros."""
    if len(historico) < 2:
        return None
    recentes = historico[-n:]
    primeiro = float(recentes[0]["peso"])
    ultimo = float(recentes[-1]["peso"])
    diff = round(ultimo - primeiro, 1)
    if diff > 0.2:
        return f"↑ +{diff}kg"
    elif diff < -0.2:
        return f"↓ {diff}kg"
    return "→ estável"


def carregar_agua_log():
    if _USA_SB:
        try:
            return _sb_req("GET", "agua_log", {"select": "*", "order": "id"}) or []
        except Exception:
            pass
    if os.path.exists(AGUA_LOG_FILE):
        with open(AGUA_LOG_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def _agua_log_local_save(lst):
    with open(AGUA_LOG_FILE, "w", encoding="utf-8") as f:
        json.dump(lst, f, ensure_ascii=False, indent=2)


def registrar_agua_log(fonte="manual"):
    agora_str = datetime.now(FUSO).strftime("%Y-%m-%d %H:%M")
    novo = {"datetime": agora_str, "fonte": fonte}
    if _USA_SB:
        try:
            _sb_req("POST", "agua_log", body=novo)
            return
        except Exception:
            pass
    lst = carregar_agua_log()
    lst.append(novo)
    _agua_log_local_save(lst)


def analisar_padroes_agua():
    """Analisa horários de confirmação de água. Retorna (contagem_por_hora, texto)."""
    logs = carregar_agua_log()
    if len(logs) < 5:
        return None, (
            f"Poucos dados ainda — só {len(logs)} confirmação(ões).\n"
            "Continua clicando em '💧 Bebi!' nos lembretes e eu monto o gráfico!"
        )

    contagem = {}
    for l in logs:
        try:
            hora = int(l["datetime"][11:13])
            contagem[hora] = contagem.get(hora, 0) + 1
        except Exception:
            continue

    if not contagem:
        return None, "Não consegui analisar os dados."

    max_c = max(contagem.values())
    total = sum(contagem.values())
    lembretes_atuais = {10, 15}

    linhas = [f"💧 *Seus horários de hidratação* ({total} confirmações):\n"]
    for hora in sorted(contagem.keys()):
        c = contagem[hora]
        barras = max(1, int((c / max_c) * 8))
        barra = "█" * barras + "░" * (8 - barras)
        tag = " ⬅ lembrete atual" if hora in lembretes_atuais else ""
        linhas.append(f"`{hora:02d}h` {barra} {c}x{tag}")

    top = sorted(contagem.keys(), key=lambda h: contagem[h], reverse=True)
    pico1 = top[0]
    pico2 = top[1] if len(top) > 1 else None

    linhas.append(f"\n🏆 Você bebe mais às *{pico1:02d}h*" +
                  (f" e *{pico2:02d}h*" if pico2 else ""))

    sugestoes = sorted([h for h in top[:2]])
    if set(sugestoes) != lembretes_atuais:
        s1, s2 = sugestoes[0], sugestoes[1] if len(sugestoes) > 1 else sugestoes[0]
        linhas.append(
            f"💡 Seus lembretes atuais são 10h e 15h, mas você hidrata mais às "
            f"{s1:02d}h e {s2:02d}h.\n"
            f"Quer ajustar? Use:\n"
            f"`/lembrar beber água` → escolha todo dia → `{s1:02d}:00`\n"
            f"`/lembrar beber água` → escolha todo dia → `{s2:02d}:00`"
        )
    else:
        linhas.append("✅ Seus lembretes estão no horário certo — continue assim!")

    return contagem, "\n".join(linhas)


def _contexto_saude():
    """Resumo de saúde para enriquecer o system prompt da IA."""
    partes = []
    historico = carregar_historico_peso()
    if historico:
        p = float(historico[-1]["peso"])
        meta = calcular_meta_agua(p)
        tend = tendencia_peso(historico)
        partes.append(f"Peso: {p}kg{(' (' + tend + ')') if tend else ''}")
        partes.append(f"Meta de água: {meta}L/dia ({int(meta/0.2)} copos de 200ml)")
    dia_letra, _ = treino_hoje()
    if dia_letra != 'Descanso':
        partes.append(f"Treino hoje: {_NOME_DIA.get(dia_letra, dia_letra)}")
    else:
        partes.append("Hoje é dia de descanso")
    streaks = carregar_streaks()
    s = streaks.get("academia", {}).get("sequencia", 0)
    if s:
        partes.append(f"Sequência academia: {s} dias")
    return "\n".join(partes)


# ── Treinos ───────────────────────────────────────────────────────────────────

def carregar_treino_dia(dia_letra):
    if _USA_SB:
        try:
            dados = _sb_req("GET", "treinos", {"dia_letra": f"eq.{dia_letra}", "select": "*", "order": "id"})
            if dados:
                return dados
        except Exception:
            pass
    return _TREINOS_FIXOS.get(dia_letra, [])


def carregar_agenda_treino():
    if _USA_SB:
        try:
            dados = _sb_req("GET", "agenda_treino", {"select": "*"})
            if dados:
                return {int(d["dia_semana"]): d["dia_letra"] for d in dados}
        except Exception:
            pass
    return _AGENDA_FIXA


def treino_hoje():
    dia_semana = datetime.now(FUSO).weekday()  # 0=segunda
    agenda = carregar_agenda_treino()
    dia_letra = agenda.get(dia_semana, 'Descanso')
    if dia_letra == 'Descanso':
        return 'Descanso', []
    return dia_letra, carregar_treino_dia(dia_letra)


def formatar_treino(dia_letra, exercicios):
    if dia_letra == 'Descanso' or not exercicios:
        return "😴 Hoje é dia de descanso! Músculo cresce no repouso — aproveita."
    linhas = [f"💪 *{_NOME_DIA.get(dia_letra, dia_letra)}*\n"]
    grupo_atual = None
    for e in exercicios:
        if e['grupo'] != grupo_atual:
            grupo_atual = e['grupo']
            icone = _ICONE_GRUPO.get(grupo_atual, '•')
            linhas.append(f"\n{icone} *{grupo_atual}*")
        linhas.append(f"  • {e['exercicio']} — {e['series']}x{e['repeticoes']}")
    return "\n".join(linhas)


# ── Tarefas ───────────────────────────────────────────────────────────────────

def carregar_tarefas():
    if _USA_SB:
        try:
            return _sb_req("GET", "tarefas", {"select": "*", "order": "id"}) or []
        except Exception:
            pass
    if os.path.exists(TAREFAS_FILE):
        with open(TAREFAS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def _tarefas_local_save(lst):
    with open(TAREFAS_FILE, "w", encoding="utf-8") as f:
        json.dump(lst, f, ensure_ascii=False, indent=2)


def adicionar_tarefa(texto, prioridade="normal", prazo=None):
    nova = {
        "texto": texto,
        "concluida": False,
        "criada": date.today().isoformat(),
        "prioridade": prioridade,
        "prazo": prazo,
    }
    if _USA_SB:
        try:
            _sb_req("POST", "tarefas", body=nova)
            return
        except Exception:
            pass
    lst = carregar_tarefas()
    lst.append(nova)
    _tarefas_local_save(lst)


def concluir_tarefa(idx):
    lst = carregar_tarefas()
    t = lst[idx - 1]
    if _USA_SB:
        try:
            _sb_req("PATCH", "tarefas", {"id": f"eq.{t['id']}"}, {"concluida": True})
            return t
        except Exception:
            pass
    lst[idx - 1]["concluida"] = True
    _tarefas_local_save(lst)
    return t


def concluir_tarefa_por_texto(descricao):
    """Conclui a primeira tarefa pendente cujo texto combine com a descrição."""
    lst = carregar_tarefas()
    alvo = descricao.lower().strip()
    palavras = [p for p in alvo.split() if len(p) > 3]
    for i, t in enumerate(lst, 1):
        if t.get("concluida"):
            continue
        texto = t.get("texto", "").lower()
        if alvo in texto or texto in alvo or any(p in texto for p in palavras):
            return concluir_tarefa(i)
    return None


def remover_tarefa(idx):
    lst = carregar_tarefas()
    t = lst[idx - 1]
    if _USA_SB:
        try:
            _sb_req("DELETE", "tarefas", {"id": f"eq.{t['id']}"})
            return t
        except Exception:
            pass
    lst.pop(idx - 1)
    _tarefas_local_save(lst)
    return t


def limpar_concluidas():
    lst = carregar_tarefas()
    concluidas = [t for t in lst if t.get("concluida")]
    if _USA_SB:
        try:
            for t in concluidas:
                _sb_req("DELETE", "tarefas", {"id": f"eq.{t['id']}"})
            return len(concluidas)
        except Exception:
            pass
    _tarefas_local_save([t for t in lst if not t.get("concluida")])
    return len(concluidas)


def formatar_tarefas(apenas_pendentes=False):
    tarefas = carregar_tarefas()
    if apenas_pendentes:
        tarefas = [t for t in tarefas if not t.get("concluida")]

    _ordem = {"alta": 0, "normal": 1, "baixa": 2}
    tarefas = sorted(
        tarefas,
        key=lambda t: (
            t.get("concluida", False),
            _ordem.get(t.get("prioridade", "normal"), 1),
            t.get("prazo") or "z",
        ),
    )

    if not tarefas:
        return "Nenhuma tarefa cadastrada."

    linhas = []
    for i, t in enumerate(tarefas, 1):
        if t.get("concluida"):
            icone = "✅"
        elif t.get("prioridade") == "alta":
            icone = "🔴"
        elif t.get("prioridade") == "baixa":
            icone = "🟢"
        else:
            icone = "⬜"
        prazo = _prazo_label(t.get("prazo")) if not t.get("concluida") else ""
        linhas.append(f"{i}. {icone} {t['texto']}{prazo}")
    return "\n".join(linhas)


# ── Notas rápidas ─────────────────────────────────────────────────────────────

def carregar_notas():
    if _USA_SB:
        try:
            return _sb_req("GET", "notas", {"select": "*", "order": "id"}) or []
        except Exception:
            pass
    if os.path.exists(NOTAS_FILE):
        with open(NOTAS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def _notas_local_save(lst):
    with open(NOTAS_FILE, "w", encoding="utf-8") as f:
        json.dump(lst, f, ensure_ascii=False, indent=2)


def adicionar_nota(texto):
    nova = {"texto": texto, "criada": datetime.now(FUSO).strftime("%Y-%m-%d %H:%M")}
    if _USA_SB:
        try:
            _sb_req("POST", "notas", body=nova)
            return
        except Exception:
            pass
    lst = carregar_notas()
    lst.append(nova)
    _notas_local_save(lst)


def remover_nota(idx):
    lst = carregar_notas()
    n = lst[idx - 1]
    if _USA_SB:
        try:
            _sb_req("DELETE", "notas", {"id": f"eq.{n['id']}"})
            return n
        except Exception:
            pass
    lst.pop(idx - 1)
    _notas_local_save(lst)
    return n


# ── Memória pessoal (fatos sobre o Bruno) ─────────────────────────────────────

def carregar_memorias():
    if _USA_SB:
        try:
            return _sb_req("GET", "memoria", {"select": "*", "order": "id"}) or []
        except Exception:
            pass
    if os.path.exists(MEMORIA_FILE):
        with open(MEMORIA_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def _memorias_local_save(lst):
    with open(MEMORIA_FILE, "w", encoding="utf-8") as f:
        json.dump(lst, f, ensure_ascii=False, indent=2)


def adicionar_memoria(fato):
    fato = fato.strip()
    if not fato:
        return
    # Evita guardar o mesmo fato duas vezes.
    if fato.lower() in [m.get("fato", "").lower() for m in carregar_memorias()]:
        return
    nova = {"fato": fato, "criada": datetime.now(FUSO).strftime("%Y-%m-%d %H:%M")}
    if _USA_SB:
        try:
            _sb_req("POST", "memoria", body=nova)
            return
        except Exception:
            pass
    lst = carregar_memorias()
    lst.append(nova)
    _memorias_local_save(lst)


def remover_memoria(idx):
    lst = carregar_memorias()
    m = lst[idx - 1]
    if _USA_SB:
        try:
            _sb_req("DELETE", "memoria", {"id": f"eq.{m['id']}"})
            return m
        except Exception:
            pass
    lst.pop(idx - 1)
    _memorias_local_save(lst)
    return m


# ── Lembretes do usuário ──────────────────────────────────────────────────────

def carregar_lembretes_usuario():
    if _USA_SB:
        try:
            return _sb_req("GET", "lembretes_usuario", {"select": "*", "order": "id"}) or []
        except Exception:
            pass
    if os.path.exists(LEMBRETES_USUARIO_FILE):
        with open(LEMBRETES_USUARIO_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def _lembretes_local_save(lst):
    with open(LEMBRETES_USUARIO_FILE, "w", encoding="utf-8") as f:
        json.dump(lst, f, ensure_ascii=False, indent=2)


def adicionar_lembrete_usuario(lembrete):
    if _USA_SB:
        try:
            _sb_req("POST", "lembretes_usuario", body=lembrete)
            return
        except Exception:
            pass
    lst = carregar_lembretes_usuario()
    lst.append(lembrete)
    _lembretes_local_save(lst)


def remover_lembrete_usuario(idx):
    lst = carregar_lembretes_usuario()
    l = lst[idx - 1]
    if _USA_SB:
        try:
            _sb_req("DELETE", "lembretes_usuario", {"id": f"eq.{l['id']}"})
            return l
        except Exception:
            pass
    lst.pop(idx - 1)
    _lembretes_local_save(lst)
    return l


def _deletar_lembrete_especifico(l):
    if _USA_SB and l.get("id"):
        try:
            _sb_req("DELETE", "lembretes_usuario", {"id": f"eq.{l['id']}"})
            return
        except Exception:
            pass
    lst = carregar_lembretes_usuario()
    lst = [x for x in lst if not (
        x["tipo"] == "especifico"
        and x.get("datetime") == l.get("datetime")
        and x.get("texto") == l.get("texto")
    )]
    _lembretes_local_save(lst)


# ── Streaks ───────────────────────────────────────────────────────────────────

def carregar_streaks():
    if _USA_SB:
        try:
            dados = _sb_req("GET", "streaks", {"select": "*"})
            return {s["habito"]: s for s in (dados or [])}
        except Exception:
            pass
    if os.path.exists(STREAKS_FILE):
        with open(STREAKS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def _streaks_local_save(streaks):
    with open(STREAKS_FILE, "w", encoding="utf-8") as f:
        json.dump(streaks, f, ensure_ascii=False, indent=2)


def confirmar_habito(habito):
    hoje = date.today().isoformat()
    ontem = (date.today() - timedelta(days=1)).isoformat()
    streaks = carregar_streaks()
    s = streaks.get(habito, {"ultima_data": None, "sequencia": 0})

    if s.get("ultima_data") == hoje:
        return s["sequencia"], None

    nova_seq = s["sequencia"] + 1 if s.get("ultima_data") == ontem else 1
    dados = {"habito": habito, "ultima_data": hoje, "sequencia": nova_seq}

    if _USA_SB:
        try:
            existente = _sb_req("GET", "streaks", {"habito": f"eq.{habito}", "select": "id"})
            if existente:
                _sb_req("PATCH", "streaks", {"habito": f"eq.{habito}"}, dados)
            else:
                _sb_req("POST", "streaks", body=dados)
        except Exception:
            pass
    else:
        streaks[habito] = dados
        _streaks_local_save(streaks)

    if nova_seq >= 30:
        msg = f"🏆 {nova_seq} dias seguidos! Você é uma lenda!"
    elif nova_seq >= 7:
        msg = f"🔥🔥 {nova_seq} dias seguidos! Incrível!"
    elif nova_seq > 1:
        msg = f"🔥 {nova_seq} dias seguidos!"
    else:
        msg = "Ótimo começo! 💪 Vamos manter!"

    return nova_seq, msg


# ── IA via Groq ───────────────────────────────────────────────────────────────

def _chamar_groq(mensagens, max_tokens=512, temperature=0.75):
    api_key = config.get("groq", {}).get("api_key", "")
    modelo = config.get("groq", {}).get("modelo", "llama-3.3-70b-versatile")
    resp = requests.post(
        "https://api.groq.com/openai/v1/chat/completions",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json={
            "model": modelo,
            "messages": mensagens,
            "max_tokens": max_tokens,
            "temperature": temperature,
        },
        timeout=30,
    )
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]


def perguntar_ia(chat_id, mensagem_usuario):
    api_key = config.get("groq", {}).get("api_key", "")
    if not api_key or "SUA_KEY" in api_key:
        return "⚠️ Configure a chave do Groq no config.json"

    # Carrega do Supabase se não estiver em memória (ex: após restart)
    if chat_id not in historico_chat:
        historico_chat[chat_id] = _hist_carregar_db(chat_id)

    agora = datetime.now(FUSO).strftime("%Y-%m-%d %H:%M")
    dia_letra, exercicios_hoje = treino_hoje()
    treino_ctx = formatar_treino(dia_letra, exercicios_hoje) if dia_letra != 'Descanso' else "Hoje é dia de descanso."

    memorias = carregar_memorias()
    if memorias:
        memoria_ctx = "\n".join(f"- {m['fato']}" for m in memorias)
    else:
        memoria_ctx = "(nada guardado ainda)"
    system_prompt = (
        "Você é o Orion, assistente pessoal do Bruno — criado por ele e totalmente dedicado a deixar a vida dele mais fácil. "
        "Você conhece o Bruno bem: trabalha na prefeitura de Saquarema, academia é prioridade, e às vezes esquece de beber água (não que você vá deixar barato).\n\n"
        "Personalidade: você é aquele amigo engraçado do grupo — faz piada, usa gíria brasileira, é sarcástico na medida certa, mas quando precisa fazer, faz rápido. "
        "Celebra quando o Bruno faz algo bom. Zoa levemente quando ele esquece as coisas. Nunca é robótico.\n\n"
        "Você também é um personal trainer informal — conhece o treino do Bruno e pode sugerir substituições de exercícios, "
        "progressão de carga, dicas de execução e ajustes baseados em dores ou limitações que ele mencionar.\n\n"
        "Regras:\n"
        "- Respostas curtas e diretas (1 a 4 linhas), com naturalidade — nada de textão\n"
        "- Varie o jeito de falar; não soe repetitivo nem robótico\n"
        "- Executa primeiro, comenta depois\n"
        "- Sem 'Olá!' nem formalidade — fala como amigo no WhatsApp\n"
        "- Emojis com moderação\n\n"
        f"Agora são {agora}.\n"
        f"O que você já sabe sobre o Bruno (use isso pra ser mais pessoal):\n{memoria_ctx}\n\n"
        f"Saúde do Bruno:\n{_contexto_saude()}\n\n"
        f"Tarefas:\n{formatar_tarefas()}\n\n"
        f"Treino de hoje:\n{treino_ctx}\n\n"
        "REGRA OBRIGATÓRIA: sempre que o Bruno pedir pra lembrar, notificar ou agendar QUALQUER coisa, "
        "coloque EXATAMENTE esta linha no início da resposta (sem espaços extras):\n"
        "[LEMBRETE:YYYY-MM-DD HH:MM:descrição]\n"
        "Use a data/hora atual para calcular horários relativos. "
        "NUNCA omita essa linha quando houver pedido de lembrete.\n"
        "Depois da linha [LEMBRETE], confirme de forma calorosa e natural que VAI lembrar, "
        "dizendo o quê e quando, pra ele saber que pode confiar "
        "(ex: 'Pode deixar! Te aviso de beber água às 13:44, não vou te deixar esquecer 💪').\n\n"
        "GESTÃO DE TAREFAS (lista do sistema):\n"
        "- Quando o Bruno pedir pra anotar/adicionar algo à lista de tarefas "
        "(ex: 'anota que preciso comprar pão', 'adiciona pagar a conta na lista'), "
        "coloque EXATAMENTE esta linha no início da resposta:\n"
        "[TAREFA:descrição curta da tarefa]\n"
        "- Quando ele disser que concluiu/terminou/fez uma tarefa "
        "(ex: 'já paguei a conta', 'terminei o relatório', 'comprei o pão'), "
        "coloque EXATAMENTE esta linha no início da resposta:\n"
        "[CONCLUIR:descrição da tarefa concluída]\n"
        "Sempre use a DESCRIÇÃO em [CONCLUIR], nunca o número. "
        "Não invente tarefas que o Bruno não pediu. Depois da linha, responda normalmente.\n\n"
        "MEMÓRIA PESSOAL: quando o Bruno contar um fato duradouro sobre ele que valha lembrar no futuro "
        "(preferências, metas, rotina, gostos, alergias, pessoas importantes, etc.), coloque no início da resposta:\n"
        "[MEMORIA:fato resumido em terceira pessoa]\n"
        "Só registre fatos realmente úteis e duradouros — ignore conversa trivial e não repita o que você já sabe.\n\n"
        "BUSCA NA WEB: quando o Bruno pedir uma informação ATUAL ou externa que você não tem certeza "
        "(preço, cotação, notícia, clima de outra cidade, resultado de jogo, fatos que mudam, ou quando ele disser 'pesquisa/procura X'), "
        "coloque EXATAMENTE esta linha no início da resposta:\n"
        "[BUSCAR:o que pesquisar]\n"
        "Prefira buscar a chutar informação que pode estar desatualizada. Use só quando precisar de info externa/atual."
    )

    historico_chat[chat_id].append({"role": "user", "content": mensagem_usuario})
    _hist_salvar(chat_id, "user", mensagem_usuario)

    mensagens = historico_chat[chat_id][-HIST_CONTEXTO:]

    try:
        resposta = _chamar_groq([{"role": "system", "content": system_prompt}] + mensagens)
        historico_chat[chat_id].append({"role": "assistant", "content": resposta})
        _hist_salvar(chat_id, "assistant", resposta)
        return resposta
    except requests.exceptions.Timeout:
        return "⏳ A IA demorou demais. Tente novamente."
    except Exception as e:
        return f"⚠️ Erro ao consultar IA: {str(e)}"


def buscar_web(termo):
    """Pesquisa na web via Tavily. Retorna texto com resumo + fontes, ou None."""
    api_key = os.environ.get("TAVILY_API_KEY", "")
    if not api_key:
        return None
    try:
        r = requests.post(
            "https://api.tavily.com/search",
            json={
                "api_key": api_key,
                "query": termo,
                "search_depth": "basic",
                "include_answer": True,
                "max_results": 4,
            },
            timeout=20,
        )
        r.raise_for_status()
        data = r.json()
        partes = []
        if data.get("answer"):
            partes.append(data["answer"])
        for res in (data.get("results") or [])[:4]:
            titulo = res.get("title", "")
            conteudo = (res.get("content", "") or "")[:300]
            if titulo or conteudo:
                partes.append(f"- {titulo}: {conteudo}")
        return "\n".join(partes) if partes else None
    except Exception:
        return None


def responder_com_busca(chat_id, termo, resultado):
    """Pede à IA pra responder usando o resultado da busca, com personalidade."""
    if not resultado:
        return "🔍 Não consegui pesquisar isso agora. Tenta de novo daqui a pouco?"
    prompt = (
        "Você é o Orion, assistente descontraído do Bruno. Ele fez uma pergunta que exigiu busca na web.\n"
        f"Termo pesquisado: '{termo}'.\n"
        f"Resultado da busca:\n{resultado}\n\n"
        "Responda ao Bruno de forma curta, natural e com sua personalidade, usando essas informações. "
        "Não invente; se a busca não trouxer a resposta, diga isso. Não use nenhuma tag tipo [BUSCAR]."
    )
    try:
        resp = _chamar_groq([{"role": "user", "content": prompt}], max_tokens=400, temperature=0.7)
        historico_chat.setdefault(chat_id, []).append({"role": "assistant", "content": resp})
        _hist_salvar(chat_id, "assistant", resp)
        return resp
    except Exception:
        return f"🔍 Achei isso:\n{resultado[:600]}"


def gerar_registro_atendimento(notas):
    """Transforma anotações informais de atendimento no texto formal do Registro."""
    api_key = config.get("groq", {}).get("api_key", "")
    if not api_key or "SUA_KEY" in api_key:
        return "⚠️ IA não configurada."
    hoje = datetime.now(FUSO).strftime("%d/%m/%Y")
    prompt = (
        "Você é o assistente do Bruno (Núcleo de Tecnologia Municipal de Rio Bonito). "
        "A partir das anotações informais dele sobre um atendimento, gere o REGISTRO DE ATENDIMENTO.\n"
        "Regras de redação:\n"
        "- Texto em 3a pessoa, impessoal e institucional (1 paragrafo, 4 a 6 frases).\n"
        "- Tom formal, sem girias e sem primeira pessoa ('eu fiz').\n"
        "- Nao repita o nome da escola no inicio da descricao.\n"
        "- Cite a data por extenso (ex: 'em 22 de junho de 2026') e os nomes com o cargo.\n"
        "- Separe o que foi RESOLVIDO (deixando explicito quando foi solucionado na hora) do que permanece PENDENTE.\n"
        f"- Use a data de hoje ({hoje}) se ele nao informar outra. NUNCA invente nomes nem datas — "
        "se faltar algo essencial, pergunte em vez de gerar.\n\n"
        "Responda EXATAMENTE neste formato (texto puro, sem markdown):\n"
        "📋 REGISTRO DE ATENDIMENTO\n\n"
        "🏫 Unidade: <nome da escola>\n"
        "📍 Modalidade: <PRESENCIAL ou ON-LINE>\n"
        "📅 Data: <dd/mm/aaaa>\n\n"
        "📝 Descrição:\n<o texto formal pronto>\n\n"
        f"Anotações do Bruno:\n{notas}"
    )
    try:
        return _chamar_groq([{"role": "user", "content": prompt}], max_tokens=700, temperature=0.4)
    except Exception as e:
        return f"⚠️ Erro ao gerar o registro: {e}"


# ── Alertas e automações ──────────────────────────────────────────────────────

def verificar_prazos():
    """Envia alerta de tarefas vencendo em até 2 dias. Chamado às 08:05."""
    chat_id = config["telegram"].get("chat_id", "")
    if not chat_id:
        return

    hoje = date.today()
    alertas = []
    for t in carregar_tarefas():
        if t.get("concluida") or not t.get("prazo"):
            continue
        try:
            prazo_dt = date.fromisoformat(t["prazo"])
            dias = (prazo_dt - hoje).days
            if dias < 0:
                alertas.append(f"⚠️ *VENCIDA:* {t['texto']} (venceu {prazo_dt.strftime('%d/%m')})")
            elif dias == 0:
                alertas.append(f"🚨 *HOJE:* {t['texto']}")
            elif dias == 1:
                alertas.append(f"⏳ *Amanhã:* {t['texto']}")
            elif dias == 2:
                alertas.append(f"📅 *Em 2 dias:* {t['texto']}")
        except Exception:
            continue

    if alertas:
        try:
            bot.send_message(chat_id, "🔔 *Atenção aos prazos:*\n\n" + "\n".join(alertas), parse_mode="Markdown")
        except Exception as e:
            print(f"Erro alerta prazos: {e}")


def enviar_resumo_semanal():
    """Resumo gerado pela IA toda sexta às 17h."""
    if datetime.now(FUSO).weekday() != 4:  # 4 = sexta-feira
        return

    chat_id = config["telegram"].get("chat_id", "")
    if not chat_id:
        return

    tarefas = carregar_tarefas()
    concluidas = [t for t in tarefas if t.get("concluida")]
    pendentes = [t for t in tarefas if not t.get("concluida")]

    try:
        nomes_c = ", ".join(t["texto"] for t in concluidas[:5]) or "nenhuma"
        nomes_p = ", ".join(t["texto"] for t in pendentes[:5]) or "nenhuma"
        prompt = (
            f"Você é o Orion. É sexta-feira, hora do resumo semanal do Bruno. "
            f"Concluídas: {len(concluidas)} ({nomes_c}). "
            f"Pendentes: {len(pendentes)} ({nomes_p}). "
            f"Gere um resumo animado em até 5 linhas: reconheça as conquistas, "
            f"destaque o que ficou pendente e motive para a próxima semana."
        )
        mensagem = _chamar_groq([{"role": "user", "content": prompt}], max_tokens=250)
    except Exception:
        mensagem = (
            f"📊 *Resumo da semana:*\n"
            f"✅ {len(concluidas)} tarefa(s) concluída(s)\n"
            f"⬜ {len(pendentes)} pendente(s)\n\n"
            "Bom fim de semana! 🎉"
        )

    try:
        bot.send_message(chat_id, mensagem, parse_mode="Markdown")
    except Exception as e:
        print(f"Erro resumo semanal: {e}")


# ── Lembretes agendados ───────────────────────────────────────────────────────

def enviar_lembrete(mensagem, habito=None):
    chat_id = config["telegram"].get("chat_id", "")
    if not chat_id:
        return
    try:
        botoes = []
        if habito == "academia":
            botoes.append(telebot.types.InlineKeyboardButton("✅ Fui!", callback_data="habito_academia"))
        elif habito == "agua":
            botoes.append(telebot.types.InlineKeyboardButton("💧 Bebi!", callback_data="habito_agua"))

        if len(_snooze_pending) > 50:
            for k in sorted(_snooze_pending.keys())[:25]:
                del _snooze_pending[k]

        snooze_id = str(int(time.time() * 1000))
        _snooze_pending[snooze_id] = mensagem
        botoes.append(telebot.types.InlineKeyboardButton("⏰ +15 min", callback_data=f"snooze_{snooze_id}"))

        markup = telebot.types.InlineKeyboardMarkup()
        markup.row(*botoes)
        bot.send_message(chat_id, mensagem, reply_markup=markup)
        print(f"[{datetime.now(FUSO).strftime('%H:%M')}] Lembrete enviado")
    except Exception as e:
        print(f"[{datetime.now(FUSO).strftime('%H:%M')}] Erro ao enviar lembrete: {e}")


def enviar_agenda_manha():
    """Bom dia personalizado com agenda do dia gerado pela IA."""
    chat_id = config["telegram"].get("chat_id", "")
    if not chat_id:
        return

    hoje = datetime.now(FUSO).strftime("%Y-%m-%d")
    tarefas = formatar_tarefas(apenas_pendentes=True)
    lembretes_hoje = [
        l for l in carregar_lembretes_usuario()
        if l["tipo"] == "especifico" and l.get("datetime", "").startswith(hoje)
    ]
    lista_lem = ", ".join(l["texto"] for l in lembretes_hoje) if lembretes_hoje else "nenhum"

    clima = buscar_clima()
    clima_str = f"Clima em Saquarema agora: {clima}." if clima else "Clima indisponível."

    dia_letra, exercicios_hoje = treino_hoje()
    treino_str = (
        formatar_treino(dia_letra, exercicios_hoje)
        if dia_letra != 'Descanso' else "Hoje é dia de descanso (sem treino)."
    )

    try:
        prompt = (
            f"Você é o Orion, assistente engraçado e descontraído do Bruno. É manhã de {hoje}. "
            f"Monte um briefing matinal com sua personalidade (pode zoar leve, usar gíria), curto e útil. "
            f"Organize de forma fluida, com quebras de linha e no máximo 1 emoji por item:\n"
            f"1) Um bom dia rápido.\n"
            f"2) {clima_str} Se previr chuva ou calor forte, avisa.\n"
            f"3) Treino de hoje: {treino_str}\n"
            f"4) Tarefas pendentes: {tarefas}. Destaque as urgentes.\n"
            f"5) Lembretes de hoje: {lista_lem}.\n"
            f"6) Feche com uma frase curta de motivação.\n"
            f"Máximo ~7 linhas no total. Não use [LEMBRETE:...], [TAREFA:...] nem [CONCLUIR:...]."
        )
        mensagem = _chamar_groq([{"role": "user", "content": prompt}], max_tokens=320, temperature=0.85)
    except Exception:
        mensagem = config.get("lembretes", [{}])[0].get(
            "mensagem", "Bom dia! ☀️\n💧 Beba água\n📋 Use /tarefas para ver o dia"
        )

    try:
        bot.send_message(chat_id, mensagem)
        print(f"[{datetime.now(FUSO).strftime('%H:%M')}] Agenda matinal enviada")
    except Exception as e:
        print(f"Erro agenda matinal: {e}")


def verificar_lembretes_especificos():
    agora = datetime.now(FUSO)
    lembretes = carregar_lembretes_usuario()
    para_disparar = [
        l for l in lembretes
        if l["tipo"] == "especifico" and (
            lambda dt: (dt.year, dt.month, dt.day, dt.hour, dt.minute) ==
                       (agora.year, agora.month, agora.day, agora.hour, agora.minute)
        )(datetime.strptime(l["datetime"], "%Y-%m-%d %H:%M"))
    ]
    for l in para_disparar:
        enviar_lembrete(f"⏰ {l['texto']}")
        _deletar_lembrete_especifico(l)


def enviar_lembrete_agua(mensagem_base):
    """Lembrete de água com meta personalizada por peso."""
    p = peso_atual()
    if p:
        meta = calcular_meta_agua(p)
        copos = int(meta / 0.2)
        mensagem = f"{mensagem_base}\n💧 Meta do dia: {meta}L (~{copos} copos)"
    else:
        mensagem = mensagem_base
    enviar_lembrete(mensagem, habito="agua")


def enviar_lembrete_academia():
    """Lembrete das 18h com o treino do dia incluído."""
    chat_id = config["telegram"].get("chat_id", "")
    if not chat_id:
        return
    dia_letra, exercicios = treino_hoje()
    treino_str = formatar_treino(dia_letra, exercicios)
    mensagem = f"🏋️ *Hora da academia!*\n\n{treino_str}"

    botoes = [telebot.types.InlineKeyboardButton("✅ Fui!", callback_data="habito_academia")]
    snooze_id = str(int(time.time() * 1000))
    _snooze_pending[snooze_id] = mensagem
    botoes.append(telebot.types.InlineKeyboardButton("⏰ +15 min", callback_data=f"snooze_{snooze_id}"))
    markup = telebot.types.InlineKeyboardMarkup()
    markup.row(*botoes)
    try:
        bot.send_message(chat_id, mensagem, reply_markup=markup, parse_mode="Markdown")
        print(f"[{datetime.now(FUSO).strftime('%H:%M')}] Lembrete academia enviado")
    except Exception as e:
        print(f"Erro lembrete academia: {e}")


def configurar_agenda():
    schedule.clear()

    schedule.every().day.at("07:00").do(enviar_agenda_manha)
    schedule.every().day.at("08:05").do(verificar_prazos)
    schedule.every().day.at("17:00").do(enviar_resumo_semanal)

    for lembrete in config.get("lembretes", []):
        if lembrete["horario"] in ("07:00",):
            continue
        if lembrete["horario"] == "18:00":
            schedule.every().day.at("18:00").do(enviar_lembrete_academia)
            continue
        if lembrete["horario"] in ("10:00", "15:00"):
            schedule.every().day.at(lembrete["horario"]).do(
                enviar_lembrete_agua, mensagem_base=lembrete["mensagem"]
            )
            continue
        habito = _HABITO_POR_HORARIO.get(lembrete["horario"])
        schedule.every().day.at(lembrete["horario"]).do(
            enviar_lembrete, mensagem=lembrete["mensagem"], habito=habito
        )

    for l in carregar_lembretes_usuario():
        if l["tipo"] == "diario":
            schedule.every().day.at(l["hora"]).do(enviar_lembrete, mensagem=f"⏰ {l['texto']}")

    schedule.every().minute.do(verificar_lembretes_especificos)
    print(f"  {len(config.get('lembretes', []))} lembretes fixos agendados.")


def thread_agenda():
    while True:
        schedule.run_pending()
        time.sleep(30)


# ── Registro de Atendimento (Diário Digital) ──────────────────────────────────

_MESES_PT = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho",
             "agosto", "setembro", "outubro", "novembro", "dezembro"]

_TEMPLATE_REGISTRO = os.path.join(os.path.dirname(__file__), "templates", "registro_atendimento.docx")


def gerar_texto_atendimento_formal(modalidade, texto_cru, data_extenso):
    """Usa a Groq para transformar anotações informais no parágrafo formal do registro."""
    system = (
        "Você redige registros oficiais de atendimento do Núcleo de Tecnologia Municipal "
        "de Rio Bonito (RJ). A partir de anotações informais, escreva UM único parágrafo "
        "formal, impessoal e em terceira pessoa (Foi realizado..., Procedeu-se..., "
        "Realizou-se...). Regras: NÃO repita o nome da escola no início; cite a data por "
        "extenso e os nomes com seus cargos quando houver; deixe explícito o que foi "
        "resolvido no próprio atendimento e o que permanece pendente; 4 a 6 frases; "
        "sem emojis, sem gírias, sem primeira pessoa. Responda APENAS com o parágrafo."
    )
    user = (
        f"Modalidade: {modalidade}\n"
        f"Data do atendimento: {data_extenso}\n"
        f"Anotações informais: {texto_cru}"
    )
    txt = _chamar_groq(
        [{"role": "system", "content": system}, {"role": "user", "content": user}],
        max_tokens=400, temperature=0.4,
    ).strip()
    return txt.strip('"').strip()


def montar_docx_atendimento(escola, modalidade, descricao_formal, hoje, out_path):
    """Preenche o template do Registro de Atendimento e salva em out_path."""
    from docx import Document
    d = Document(_TEMPLATE_REGISTRO)

    # [2] Unidade Escolar — nome sublinhado, mantendo as linhas
    p2 = d.paragraphs[2]
    r0 = p2.runs[0]
    label = "Unidade Escolar: "
    unders = r0.text[len(label):]
    r0.text = label
    rn = p2.add_run(escola)
    rn.underline = True
    p2.add_run(unders[len(escola):] if len(unders) > len(escola) else unders)

    # [4] marca X em ON-LINE (run4) ou PRESENCIAL (run2)
    alvo = 4 if modalidade.upper().startswith("ON") else 2
    d.paragraphs[4].runs[alvo].text = "( X"

    # [6] descrição formal sublinhada, sem linhas em branco extras
    p6 = d.paragraphs[6]
    for r in list(p6.runs):
        r.text = ""
    p6.runs[0].text = descricao_formal
    p6.runs[0].underline = True

    # [12] data de hoje
    p12 = d.paragraphs[12]
    p12.runs[0].text = f"Rio Bonito {hoje.day:02d} / {hoje.month:02d} / {hoje.year}"
    for r in p12.runs[1:]:
        r.text = ""

    d.save(out_path)
    return out_path


# ── Handlers ──────────────────────────────────────────────────────────────────

def registrar_handlers():

    @bot.message_handler(commands=["start", "ajuda"])
    def cmd_start(msg):
        chat_id_str = str(msg.chat.id)
        if not config["telegram"].get("chat_id"):
            config["telegram"]["chat_id"] = chat_id_str
            salvar_config()
            extra = "✅ *Seu chat foi vinculado! Você receberá os lembretes aqui.*\n\n"
        else:
            extra = ""
        bot.reply_to(msg, (
            f"{extra}"
            "🌟 *Orion — Assistente Pessoal*\n\n"
            "📋 *Tarefas:*\n"
            "/adicionar [texto] [@DD/MM] — adicionar (use ! para urgente)\n"
            "/urgente [texto] — adicionar tarefa urgente\n"
            "/tarefas — ver tarefas (ordenadas por prioridade)\n"
            "/concluir [n] — marcar como feita\n"
            "/remover [n] — remover\n"
            "/historico — ver tarefas concluídas\n\n"
            "📝 *Notas rápidas:*\n"
            "/nota [texto] — salvar nota\n"
            "/notas — ver todas as notas\n"
            "/deletar\\_nota [n] — remover nota\n\n"
            "⏰ *Lembretes:*\n"
            "/lembrar [texto] — criar lembrete\n"
            "/timer [min] [descrição] — timer rápido\n"
            "/meus\\_lembretes — ver lembretes criados\n"
            "/cancelar\\_lembrete [n] — remover lembrete\n\n"
            "💪 *Saúde:*\n"
            "/peso 80.5 — registrar peso\n"
            "/historico\\_peso — histórico de peso com tendência\n"
            "/meta\\_agua — ver meta de água pelo seu peso\n"
            "/stats — dashboard completo (peso, treino, hábitos, tarefas)\n"
            "/dicas — dicas personalizadas da IA\n"
            "/bebi — registrar copo de água agora\n"
            "/analise\\_agua — ver gráfico de quando você bebe água\n\n"
            "🏋️ *Academia:*\n"
            "/treino — treino de hoje\n"
            "/treino A — ver Dia A (Peito/Tríceps/Ombro)\n"
            "/treino B — ver Dia B (Pernas)\n"
            "/treino C — ver Dia C (Costas/Bíceps)\n"
            "/add\\_exercicio A Peito Supino 4 6-8 — adicionar exercício\n\n"
            "📄 *Documentos:*\n"
            "/registro Escola | online | o que foi feito — gera o Registro de Atendimento\n\n"
            "🔍 *Busca:*\n"
            "/buscar [termo] — busca em tarefas e notas\n\n"
            "🏆 *Hábitos:*\n"
            "/streaks — ver sequência de hábitos\n\n"
            "💬 *Ou manda uma mensagem — o Orion responde!*\n"
            "_Ex: 'me lembra da reunião amanhã às 14h'_"
        ), parse_mode="Markdown")

    # ── Registro de Atendimento ──

    @bot.message_handler(commands=["registro", "atendimento"])
    def cmd_registro(msg):
        partes_cmd = msg.text.split(None, 1)
        raw = partes_cmd[1].strip() if len(partes_cmd) > 1 else ""
        partes = [p.strip() for p in raw.split("|")]
        if len(partes) < 3 or not all(partes[:2]) or not partes[2]:
            bot.reply_to(msg, (
                "⚠️ *Uso:* `/registro Escola | online ou presencial | o que foi feito`\n\n"
                "*Ex:* `/registro E.M. Maurício Kopke | online | resolvi o acesso aos "
                "lançamentos de notas de uma professora na hora, corrigi frequências "
                "incorretas de uma turma, falta carga horária da ATA e ficha individual`"
            ), parse_mode="Markdown")
            return

        escola = partes[0]
        modalidade = "ON-LINE" if partes[1].lower().startswith("on") else "PRESENCIAL"
        texto_cru = "|".join(partes[2:]).strip()

        bot.send_chat_action(msg.chat.id, "typing")
        hoje = datetime.now(FUSO)
        data_ext = f"{hoje.day} de {_MESES_PT[hoje.month - 1]} de {hoje.year}"

        try:
            descricao = gerar_texto_atendimento_formal(modalidade, texto_cru, data_ext)
        except Exception as e:
            bot.reply_to(msg, f"⚠️ Erro ao gerar o texto pela IA: {e}")
            return

        out = os.path.join(
            tempfile.gettempdir(),
            f"registro_atendimento_{hoje:%Y%m%d_%H%M%S}.docx",
        )
        try:
            montar_docx_atendimento(escola, modalidade, descricao, hoje, out)
        except Exception as e:
            bot.reply_to(msg, f"⚠️ Texto gerado, mas falhou montar o documento: {e}\n\n{descricao}")
            return

        try:
            with open(out, "rb") as f:
                bot.send_document(
                    msg.chat.id, f,
                    visible_file_name=f"Registro de Atendimento - {escola}.docx",
                    caption=f"📄 Registro de Atendimento — {escola} ({modalidade}) — {hoje:%d/%m/%Y}",
                )
            bot.send_message(msg.chat.id, f"📝 *Texto gerado:*\n\n{descricao}", parse_mode="Markdown")
        finally:
            try:
                os.remove(out)
            except OSError:
                pass

    # ── Tarefas ──

    @bot.message_handler(commands=["tarefas"])
    def cmd_tarefas(msg):
        bot.reply_to(msg, f"📋 *Suas tarefas:*\n\n{formatar_tarefas()}", parse_mode="Markdown")

    @bot.message_handler(commands=["adicionar"])
    def cmd_adicionar(msg):
        texto_raw = msg.text.replace("/adicionar", "").strip()
        if not texto_raw:
            bot.reply_to(msg,
                         "⚠️ Uso: `/adicionar texto [@DD/MM]`\n"
                         "Use `!` no início para urgente.\n"
                         "Ex: `/adicionar ! Relatório @30/06`",
                         parse_mode="Markdown")
            return
        texto, prioridade, prazo = _parse_tarefa(texto_raw)
        adicionar_tarefa(texto, prioridade, prazo)
        prazo_str = f"\n📅 Prazo: {date.fromisoformat(prazo).strftime('%d/%m/%Y')}" if prazo else ""
        prioridade_str = " 🔴 *URGENTE*" if prioridade == "alta" else ""
        bot.reply_to(msg, f"✅ Tarefa adicionada:{prioridade_str}\n*{texto}*{prazo_str}", parse_mode="Markdown")

    @bot.message_handler(commands=["urgente"])
    def cmd_urgente(msg):
        texto_raw = msg.text.replace("/urgente", "").strip()
        if not texto_raw:
            bot.reply_to(msg, "⚠️ Uso: `/urgente Descrição da tarefa`", parse_mode="Markdown")
            return
        texto, _, prazo = _parse_tarefa(texto_raw)
        adicionar_tarefa(texto, prioridade="alta", prazo=prazo)
        prazo_str = f"\n📅 Prazo: {date.fromisoformat(prazo).strftime('%d/%m/%Y')}" if prazo else ""
        bot.reply_to(msg, f"🔴 *Tarefa urgente adicionada:*\n{texto}{prazo_str}", parse_mode="Markdown")

    @bot.message_handler(commands=["concluir"])
    def cmd_concluir(msg):
        try:
            n = int(msg.text.replace("/concluir", "").strip())
            t = concluir_tarefa(n)
            bot.reply_to(msg, f"✅ Tarefa {n} concluída: *{t['texto']}*", parse_mode="Markdown")
        except (ValueError, IndexError):
            bot.reply_to(msg, "⚠️ Uso: `/concluir 1` (número da tarefa)", parse_mode="Markdown")

    @bot.message_handler(commands=["remover"])
    def cmd_remover(msg):
        try:
            n = int(msg.text.replace("/remover", "").strip())
            t = remover_tarefa(n)
            bot.reply_to(msg, f"🗑️ Tarefa removida: *{t['texto']}*", parse_mode="Markdown")
        except (ValueError, IndexError):
            bot.reply_to(msg, "⚠️ Uso: `/remover 1` (número da tarefa)", parse_mode="Markdown")

    @bot.message_handler(commands=["limpar"])
    def cmd_limpar(msg):
        n = limpar_concluidas()
        bot.reply_to(msg, f"🧹 {n} tarefa(s) concluída(s) removida(s).")

    @bot.message_handler(commands=["historico"])
    def cmd_historico(msg):
        concluidas = [t for t in carregar_tarefas() if t.get("concluida")]
        if not concluidas:
            bot.reply_to(msg, "Nenhuma tarefa concluída ainda. Bora trabalhar! 💪")
            return
        linhas = [f"{i}. ✅ {t['texto']}" for i, t in enumerate(concluidas, 1)]
        bot.reply_to(msg, "✅ *Tarefas concluídas:*\n\n" + "\n".join(linhas), parse_mode="Markdown")

    # ── Notas ──

    @bot.message_handler(commands=["nota"])
    def cmd_nota(msg):
        texto = msg.text.replace("/nota", "").strip()
        if not texto:
            bot.reply_to(msg, "⚠️ Uso: `/nota Texto da nota`\nEx: `/nota Protocolo 2024/001 - aguardando sec`", parse_mode="Markdown")
            return
        adicionar_nota(texto)
        bot.reply_to(msg, f"📝 Nota salva:\n_{texto}_", parse_mode="Markdown")

    @bot.message_handler(commands=["notas"])
    def cmd_notas(msg):
        notas = carregar_notas()
        if not notas:
            bot.reply_to(msg, "Nenhuma nota salva.\nUse `/nota texto` para adicionar.", parse_mode="Markdown")
            return
        linhas = [f"{i}. 📝 {n['texto']}" for i, n in enumerate(notas, 1)]
        bot.reply_to(msg, "📝 *Suas notas:*\n\n" + "\n".join(linhas) +
                     "\n\nUse `/deletar_nota [número]` para remover.", parse_mode="Markdown")

    @bot.message_handler(commands=["deletar_nota"])
    def cmd_deletar_nota(msg):
        try:
            n = int(msg.text.replace("/deletar_nota", "").strip())
            nota = remover_nota(n)
            bot.reply_to(msg, f"🗑️ Nota removida: _{nota['texto']}_", parse_mode="Markdown")
        except (ValueError, IndexError):
            bot.reply_to(msg, "⚠️ Uso: `/deletar_nota 1` (número da nota)", parse_mode="Markdown")

    # ── Memória pessoal ──

    @bot.message_handler(commands=["memorias"])
    def cmd_memorias(msg):
        memorias = carregar_memorias()
        if not memorias:
            bot.reply_to(msg, (
                "🧠 Ainda não guardei nada sobre você.\n"
                "É só me contar coisas no chat (ex: 'minha meta é chegar a 80kg') "
                "que eu vou lembrando sozinho."
            ))
            return
        linhas = [f"{i}. 🧠 {m['fato']}" for i, m in enumerate(memorias, 1)]
        bot.reply_to(msg, "🧠 *O que eu sei sobre você:*\n\n" + "\n".join(linhas) +
                     "\n\nUse `/esquecer [número]` pra eu apagar.", parse_mode="Markdown")

    @bot.message_handler(commands=["esquecer"])
    def cmd_esquecer(msg):
        try:
            n = int(msg.text.replace("/esquecer", "").strip())
            m = remover_memoria(n)
            bot.reply_to(msg, f"🗑️ Esqueci: _{m['fato']}_", parse_mode="Markdown")
        except (ValueError, IndexError):
            bot.reply_to(msg, "⚠️ Uso: `/esquecer 1` (número da memória, veja em /memorias)", parse_mode="Markdown")

    # ── Streaks ──

    @bot.message_handler(commands=["streaks"])
    def cmd_streaks(msg):
        streaks = carregar_streaks()
        if not streaks:
            bot.reply_to(msg, (
                "Nenhum hábito registrado ainda.\n"
                "Quando o Orion mandar um lembrete de água ou academia, "
                "clique no botão para confirmar e a sequência começa! 💪"
            ))
            return
        nomes = {"academia": "🏋️ Academia", "agua": "💧 Hidratação"}
        linhas = []
        for habito, s in streaks.items():
            nome = nomes.get(habito, habito)
            seq = s.get("sequencia", 0)
            emoji = "🔥" if seq >= 3 else ("✅" if seq > 0 else "⬜")
            linhas.append(f"{emoji} {nome}: *{seq} dia(s) seguido(s)*")
        bot.reply_to(msg, "🏆 *Seus hábitos:*\n\n" + "\n".join(linhas), parse_mode="Markdown")

    # ── Lembretes ──

    @bot.message_handler(commands=["lembrar"])
    def cmd_lembrar(msg):
        texto = msg.text.replace("/lembrar", "").strip()
        if not texto:
            bot.reply_to(msg, "⚠️ Uso: `/lembrar Descrição do lembrete`", parse_mode="Markdown")
            return
        estados[msg.chat.id] = {"texto": texto, "step": "tipo"}
        markup = telebot.types.InlineKeyboardMarkup()
        markup.row(
            telebot.types.InlineKeyboardButton("📅 Data específica", callback_data="lembrar_especifico"),
            telebot.types.InlineKeyboardButton("🔁 Todo dia", callback_data="lembrar_diario"),
        )
        bot.reply_to(msg, f'📝 *"{texto}"*\n\nEste lembrete é:', reply_markup=markup, parse_mode="Markdown")

    @bot.callback_query_handler(func=lambda c: c.data in ("lembrar_especifico", "lembrar_diario"))
    def callback_tipo_lembrete(call):
        chat_id = call.message.chat.id
        if chat_id not in estados:
            bot.answer_callback_query(call.id, "Sessão expirada. Use /lembrar novamente.")
            return
        if call.data == "lembrar_especifico":
            estados[chat_id]["tipo"] = "especifico"
            estados[chat_id]["step"] = "data"
            bot.edit_message_text(
                f'📝 *"{estados[chat_id]["texto"]}"*\n\n📅 Qual data e hora?\n'
                'Formato: `DD/MM HH:MM`\nExemplo: `25/06 14:30` ou `hoje 15:00`',
                chat_id=chat_id, message_id=call.message.message_id, parse_mode="Markdown"
            )
        else:
            estados[chat_id]["tipo"] = "diario"
            estados[chat_id]["step"] = "hora"
            bot.edit_message_text(
                f'📝 *"{estados[chat_id]["texto"]}"*\n\n⏰ Que horas todo dia?\n'
                'Formato: `HH:MM`\nExemplo: `08:30`',
                chat_id=chat_id, message_id=call.message.message_id, parse_mode="Markdown"
            )
        bot.answer_callback_query(call.id)

    @bot.callback_query_handler(func=lambda c: c.data.startswith("snooze_"))
    def callback_snooze(call):
        snooze_id = call.data[len("snooze_"):]
        mensagem = _snooze_pending.pop(snooze_id, None)
        if not mensagem:
            bot.answer_callback_query(call.id, "⚠️ Lembrete expirado.")
            return
        dt_snooze = datetime.now(FUSO) + timedelta(minutes=15)
        texto = re.sub(r"^⏰\s*", "", mensagem).strip()
        adicionar_lembrete_usuario({
            "tipo": "especifico",
            "texto": texto,
            "datetime": dt_snooze.strftime("%Y-%m-%d %H:%M"),
        })
        bot.answer_callback_query(call.id, "⏰ Adiado 15 min!")
        try:
            bot.edit_message_reply_markup(call.message.chat.id, call.message.message_id, reply_markup=None)
        except Exception:
            pass
        bot.send_message(call.message.chat.id,
                         f"⏰ Ok! Vou te lembrar às *{dt_snooze.strftime('%H:%M')}*.",
                         parse_mode="Markdown")

    @bot.callback_query_handler(func=lambda c: c.data.startswith("habito_"))
    def callback_habito(call):
        habito = call.data[len("habito_"):]
        seq, msg_streak = confirmar_habito(habito)

        if habito == "agua":
            registrar_agua_log(fonte="lembrete")

        if msg_streak is None:
            bot.answer_callback_query(call.id, "Já confirmado hoje! 👍")
            return
        nomes = {"academia": "Academia", "agua": "Hidratação"}
        bot.answer_callback_query(call.id, msg_streak)
        try:
            bot.edit_message_reply_markup(call.message.chat.id, call.message.message_id, reply_markup=None)
        except Exception:
            pass
        bot.send_message(call.message.chat.id,
                         f"✅ *{nomes.get(habito, habito)}* confirmada!\n{msg_streak}",
                         parse_mode="Markdown")

    @bot.message_handler(commands=["meus_lembretes"])
    def cmd_meus_lembretes(msg):
        lembretes = carregar_lembretes_usuario()
        if not lembretes:
            bot.reply_to(msg, "Nenhum lembrete personalizado.\nUse /lembrar para adicionar!")
            return
        linhas = []
        for i, l in enumerate(lembretes, 1):
            if l["tipo"] == "especifico":
                dt = datetime.strptime(l["datetime"], "%Y-%m-%d %H:%M")
                linhas.append(f"{i}. 📅 {dt.strftime('%d/%m às %H:%M')} — {l['texto']}")
            else:
                linhas.append(f"{i}. 🔁 Todo dia às {l['hora']} — {l['texto']}")
        bot.reply_to(msg, "📋 *Seus lembretes:*\n\n" + "\n".join(linhas) +
                     "\n\nUse `/cancelar_lembrete [número]` para remover.", parse_mode="Markdown")

    @bot.message_handler(commands=["cancelar_lembrete"])
    def cmd_cancelar_lembrete(msg):
        try:
            n = int(msg.text.replace("/cancelar_lembrete", "").strip())
            l = remover_lembrete_usuario(n)
            configurar_agenda()
            bot.reply_to(msg, f"🗑️ Lembrete removido: *{l['texto']}*", parse_mode="Markdown")
        except (ValueError, IndexError):
            bot.reply_to(msg, "⚠️ Uso: `/cancelar_lembrete 1`", parse_mode="Markdown")

    @bot.message_handler(commands=["lembretes"])
    def cmd_lembretes(msg):
        lembretes = config.get("lembretes", [])
        if not lembretes:
            bot.reply_to(msg, "Nenhum lembrete fixo configurado.")
            return
        linhas = [f"⏰ `{l['horario']}` — {l['mensagem'].splitlines()[0]}" for l in lembretes]
        bot.reply_to(msg, "📅 *Lembretes fixos:*\n\n" + "\n".join(linhas), parse_mode="Markdown")

    @bot.message_handler(commands=["timer"])
    def cmd_timer(msg):
        partes = msg.text.replace("/timer", "").strip().split(None, 1)
        if not partes:
            bot.reply_to(msg, "⚠️ Uso: `/timer 25` ou `/timer 30 Ligar pro cliente`", parse_mode="Markdown")
            return
        try:
            minutos = int(partes[0])
            descricao = partes[1] if len(partes) > 1 else "Timer finalizado!"
            dt = datetime.now(FUSO) + timedelta(minutes=minutos)
            adicionar_lembrete_usuario({
                "tipo": "especifico",
                "texto": f"⏱️ {descricao}",
                "datetime": dt.strftime("%Y-%m-%d %H:%M"),
            })
            bot.reply_to(msg,
                         f"⏱️ Timer de *{minutos} min* configurado!\nVou te chamar às *{dt.strftime('%H:%M')}*.",
                         parse_mode="Markdown")
        except ValueError:
            bot.reply_to(msg, "⚠️ Uso: `/timer 25` ou `/timer 30 Ligar pro cliente`", parse_mode="Markdown")

    @bot.message_handler(commands=["peso"])
    def cmd_peso(msg):
        partes = msg.text.replace("/peso", "").strip().split(None, 1)
        if not partes:
            bot.reply_to(msg, "⚠️ Uso: `/peso 80.5` ou `/peso 80.5 pós-treino`", parse_mode="Markdown")
            return
        try:
            p = float(partes[0].replace(",", "."))
            obs = partes[1] if len(partes) > 1 else None
            registrar_peso(p, obs)
            meta = calcular_meta_agua(p)
            copos = int(meta / 0.2)
            historico = carregar_historico_peso()
            tend = tendencia_peso(historico)
            tend_str = f" ({tend})" if tend else ""
            bot.reply_to(msg,
                f"⚖️ *{p}kg registrado!*{tend_str}\n"
                f"💧 Sua meta de água: *{meta}L/dia* (~{copos} copos)",
                parse_mode="Markdown")
        except ValueError:
            bot.reply_to(msg, "⚠️ Uso: `/peso 80.5`", parse_mode="Markdown")

    @bot.message_handler(commands=["historico_peso"])
    def cmd_historico_peso(msg):
        historico = carregar_historico_peso()
        if not historico:
            bot.reply_to(msg, "Nenhum peso registrado ainda.\nUse `/peso 80.5` para começar.", parse_mode="Markdown")
            return
        recentes = historico[-10:]
        linhas = ["📈 *Histórico de peso:*\n"]
        for r in recentes:
            data_fmt = datetime.strptime(r["data"], "%Y-%m-%d").strftime("%d/%m")
            obs = f" — _{r['observacao']}_" if r.get("observacao") else ""
            linhas.append(f"  `{data_fmt}` {r['peso']}kg{obs}")
        tend = tendencia_peso(historico)
        if tend:
            linhas.append(f"\nTendência: *{tend}*")
        bot.reply_to(msg, "\n".join(linhas), parse_mode="Markdown")

    @bot.message_handler(commands=["stats"])
    def cmd_stats(msg):
        historico = carregar_historico_peso()
        streaks = carregar_streaks()
        tarefas = carregar_tarefas()
        hoje = date.today().isoformat()

        linhas = ["📊 *Suas estatísticas:*\n"]

        # Peso
        if historico:
            p = float(historico[-1]["peso"])
            tend = tendencia_peso(historico)
            tend_str = f" ({tend})" if tend else ""
            meta = calcular_meta_agua(p)
            linhas.append(f"⚖️ Peso: *{p}kg*{tend_str}")
            linhas.append(f"💧 Meta de água: *{meta}L/dia* (~{int(meta/0.2)} copos)")
        else:
            linhas.append("⚖️ Peso: _não registrado — use /peso_")

        # Treino de hoje
        dia_letra, _ = treino_hoje()
        linhas.append(f"🏋️ Hoje: *{_NOME_DIA.get(dia_letra, 'Descanso')}*")

        # Streaks
        s_ac = streaks.get("academia", {}).get("sequencia", 0)
        s_ag = streaks.get("agua", {}).get("sequencia", 0)
        linhas.append(f"🔥 Sequência academia: *{s_ac} dia(s)*")
        linhas.append(f"💧 Sequência hidratação: *{s_ag} dia(s)*")

        # Tarefas
        pendentes = [t for t in tarefas if not t.get("concluida")]
        urgentes = [t for t in pendentes if t.get("prioridade") == "alta"]
        linhas.append(f"📋 Tarefas pendentes: *{len(pendentes)}*" +
                      (f" ({len(urgentes)} urgentes 🔴)" if urgentes else ""))

        bot.reply_to(msg, "\n".join(linhas), parse_mode="Markdown")

    @bot.message_handler(commands=["dicas"])
    def cmd_dicas(msg):
        api_key = config.get("groq", {}).get("api_key", "")
        if not api_key or "SUA_KEY" in api_key:
            bot.reply_to(msg, "⚠️ IA não configurada.")
            return
        bot.send_chat_action(msg.chat.id, "typing")
        historico = carregar_historico_peso()
        streaks = carregar_streaks()
        dia_letra, exercicios = treino_hoje()
        p = float(historico[-1]["peso"]) if historico else None
        tend = tendencia_peso(historico) if historico else None
        s_ac = streaks.get("academia", {}).get("sequencia", 0)

        prompt = (
            f"Você é o Orion, personal trainer e assistente do Bruno. "
            f"Dados: peso {p}kg, tendência {tend}, sequência academia {s_ac} dias. "
            f"Treino de hoje: {_NOME_DIA.get(dia_letra, 'Descanso')}. "
            f"Dê 3 dicas práticas e personalizadas de treino, nutrição ou recuperação. "
            f"Seja direto, use linguagem informal, máximo 6 linhas no total."
        )
        try:
            resposta = _chamar_groq([{"role": "user", "content": prompt}], max_tokens=300)
            bot.reply_to(msg, resposta)
        except Exception as e:
            bot.reply_to(msg, f"⚠️ Erro: {e}")

    @bot.message_handler(commands=["meta_agua"])
    def cmd_meta_agua(msg):
        p = peso_atual()
        if not p:
            bot.reply_to(msg,
                "Ainda não registrei seu peso.\nUse `/peso 80.5` primeiro!",
                parse_mode="Markdown")
            return
        meta = calcular_meta_agua(p)
        copos = int(meta / 0.2)
        bot.reply_to(msg,
            f"💧 *Meta de água personalizada:*\n\n"
            f"Baseado no seu peso de *{p}kg*:\n"
            f"→ *{meta}L por dia* (~{copos} copos de 200ml)\n\n"
            f"_Fórmula: 35ml × peso corporal_",
            parse_mode="Markdown")

    @bot.message_handler(commands=["limpar_historico"])
    def cmd_limpar_historico(msg):
        chat_id = msg.chat.id
        historico_chat.pop(chat_id, None)
        _hist_limpar_db(chat_id)
        bot.reply_to(msg, "🧹 Histórico de conversa limpo!\nComeçamos do zero.")

    @bot.message_handler(commands=["bebi"])
    def cmd_bebi(msg):
        registrar_agua_log(fonte="manual")
        confirmar_habito("agua")
        agora_str = datetime.now(FUSO).strftime("%H:%M")
        logs = carregar_agua_log()
        hoje = date.today().isoformat()
        hoje_count = sum(1 for l in logs if l["datetime"].startswith(hoje))
        p = peso_atual()
        meta_str = ""
        if p:
            meta = calcular_meta_agua(p)
            copos = int(meta / 0.2)
            meta_str = f"\nHoje: *{hoje_count}/{copos} copos* (meta {meta}L)"
        bot.reply_to(msg, f"💧 Água às *{agora_str}* registrada!{meta_str}", parse_mode="Markdown")

    @bot.message_handler(commands=["analise_agua"])
    def cmd_analise_agua(msg):
        bot.send_chat_action(msg.chat.id, "typing")
        _, texto = analisar_padroes_agua()
        bot.reply_to(msg, texto, parse_mode="Markdown")

    @bot.message_handler(commands=["treino"])
    def cmd_treino(msg):
        arg = msg.text.replace("/treino", "").strip().upper()
        if arg in ('A', 'B', 'C'):
            exercicios = carregar_treino_dia(arg)
            bot.reply_to(msg, formatar_treino(arg, exercicios), parse_mode="Markdown")
        else:
            dia_letra, exercicios = treino_hoje()
            dias_semana = ['segunda', 'terça', 'quarta', 'quinta', 'sexta', 'sábado', 'domingo']
            hoje_nome = dias_semana[datetime.now(FUSO).weekday()]
            texto = formatar_treino(dia_letra, exercicios)
            bot.reply_to(msg, f"📅 *{hoje_nome.capitalize()}:*\n{texto}", parse_mode="Markdown")

    @bot.message_handler(commands=["add_exercicio"])
    def cmd_add_exercicio(msg):
        partes = msg.text.replace("/add_exercicio", "").strip().split(None, 3)
        if len(partes) < 4:
            bot.reply_to(msg,
                "⚠️ Uso: `/add_exercicio A Peito Supino reto 4 6-8`\n"
                "Formato: `/add_exercicio [A/B/C] [Grupo] [Exercício] [Séries] [Reps]`",
                parse_mode="Markdown")
            return
        try:
            dia_letra = partes[0].upper()
            grupo = partes[1]
            resto = partes[2].rsplit(None, 2)
            if len(resto) < 3:
                raise ValueError
            exercicio, series, repeticoes = resto
            if dia_letra not in ('A', 'B', 'C'):
                raise ValueError
            novo = {"dia_letra": dia_letra, "grupo": grupo, "exercicio": exercicio,
                    "series": series, "repeticoes": repeticoes}
            if _USA_SB:
                _sb_req("POST", "treinos", body=novo)
            bot.reply_to(msg,
                f"✅ Exercício adicionado ao *Dia {dia_letra}*!\n"
                f"  {grupo} — {exercicio} {series}x{repeticoes}",
                parse_mode="Markdown")
        except Exception:
            bot.reply_to(msg,
                "⚠️ Formato: `/add_exercicio A Peito Supino reto 4 6-8`",
                parse_mode="Markdown")

    @bot.message_handler(commands=["buscar"])
    def cmd_buscar(msg):
        termo = msg.text.replace("/buscar", "").strip().lower()
        if not termo:
            bot.reply_to(msg, "⚠️ Uso: `/buscar protocolo`", parse_mode="Markdown")
            return
        tarefas_match = [t for t in carregar_tarefas() if termo in t["texto"].lower()]
        notas_match = [n for n in carregar_notas() if termo in n["texto"].lower()]
        if not tarefas_match and not notas_match:
            bot.reply_to(msg, f"🔍 Nada encontrado para *{termo}*.", parse_mode="Markdown")
            return
        linhas = [f"🔍 Resultados para *{termo}*:\n"]
        if tarefas_match:
            linhas.append("📋 *Tarefas:*")
            for t in tarefas_match:
                icone = "✅" if t.get("concluida") else ("🔴" if t.get("prioridade") == "alta" else "⬜")
                prazo = _prazo_label(t.get("prazo")) if not t.get("concluida") else ""
                linhas.append(f"  {icone} {t['texto']}{prazo}")
        if notas_match:
            linhas.append("\n📝 *Notas:*")
            for n in notas_match:
                linhas.append(f"  📝 {n['texto']}")
        bot.reply_to(msg, "\n".join(linhas), parse_mode="Markdown")

    @bot.message_handler(commands=["registro"])
    def cmd_registro(msg):
        notas = msg.text.replace("/registro", "", 1).strip()
        if notas:
            bot.send_chat_action(msg.chat.id, "typing")
            bot.reply_to(msg, gerar_registro_atendimento(notas))
        else:
            estados[msg.chat.id] = {"step": "registro"}
            bot.reply_to(msg, (
                "📋 Registro de Atendimento\n\n"
                "Me manda as anotações do atendimento (pode ser bem informal), tipo:\n"
                '"escola Posse, online, secretário Jorge e secretária Ana, resolvi o acesso a notas '
                'de uma professora, corrigi frequências de uma turma, falta carga horária da ATA"\n\n'
                "Que eu te devolvo o texto formal pronto. 👇"
            ))

    # ── Chat livre ──

    @bot.message_handler(func=lambda m: True, content_types=["text"])
    def handle_texto(msg):
        chat_id = msg.chat.id
        # Vincula o chat em qualquer mensagem (não só no /start), pois no Railway
        # o config é efêmero e o chat_id é perdido a cada restart.
        if not config["telegram"].get("chat_id"):
            config["telegram"]["chat_id"] = str(chat_id)
            salvar_config()
        estado = estados.get(chat_id, {})

        if estado.get("step") == "data":
            try:
                partes = msg.text.strip().split(" ")
                hora_str = partes[-1]
                dia_mes = partes[0].lower()
                if dia_mes == "hoje":
                    d = date.today()
                else:
                    dia, mes = dia_mes.split("/")
                    d = date(date.today().year, int(mes), int(dia))
                hora, minuto = hora_str.split(":")
                dt = datetime(d.year, d.month, d.day, int(hora), int(minuto))
                adicionar_lembrete_usuario({
                    "tipo": "especifico",
                    "texto": estado["texto"],
                    "datetime": dt.strftime("%Y-%m-%d %H:%M"),
                })
                del estados[chat_id]
                bot.reply_to(msg,
                             f"✅ Lembrete agendado!\n📅 *{dt.strftime('%d/%m/%Y às %H:%M')}*\n📝 {estado['texto']}",
                             parse_mode="Markdown")
            except Exception:
                bot.reply_to(msg,
                             "⚠️ Formato inválido. Use: `DD/MM HH:MM` ou `hoje HH:MM`",
                             parse_mode="Markdown")

        elif estado.get("step") == "hora":
            try:
                hora_str = msg.text.strip()
                h, m = hora_str.split(":")
                int(h); int(m)
                adicionar_lembrete_usuario({"tipo": "diario", "texto": estado["texto"], "hora": hora_str})
                configurar_agenda()
                del estados[chat_id]
                bot.reply_to(msg,
                             f"✅ Lembrete diário criado!\n🔁 *Todo dia às {hora_str}*\n📝 {estado['texto']}",
                             parse_mode="Markdown")
            except Exception:
                bot.reply_to(msg, "⚠️ Formato inválido. Use: `HH:MM`", parse_mode="Markdown")

        elif estado.get("step") == "registro":
            del estados[chat_id]
            bot.send_chat_action(chat_id, "typing")
            bot.reply_to(msg, gerar_registro_atendimento(msg.text))

        else:
            bot.send_chat_action(chat_id, "typing")
            resposta = perguntar_ia(chat_id, msg.text)

            m_buscar = re.search(r'\[BUSCAR\s*:\s*(.+?)\]', resposta)
            if m_buscar:
                termo = m_buscar.group(1).strip()
                bot.send_chat_action(chat_id, "typing")
                resultado = buscar_web(termo)
                resposta = responder_com_busca(chat_id, termo, resultado)

            match = re.search(r'\[LEMBRETE\s*:\s*(\d{4}-\d{2}-\d{2})\s+(\d{2}:\d{2})\s*:\s*(.+?)\]', resposta)
            if match:
                try:
                    dt = datetime.strptime(f"{match.group(1)} {match.group(2)}", "%Y-%m-%d %H:%M")
                    texto_lembrete = match.group(3).strip()
                    adicionar_lembrete_usuario({
                        "tipo": "especifico",
                        "texto": texto_lembrete,
                        "datetime": dt.strftime("%Y-%m-%d %H:%M"),
                    })
                    resposta = re.sub(r'\[LEMBRETE[^\]]+\]', '', resposta).strip()
                except Exception:
                    pass

            m_tarefa = re.search(r'\[TAREFA\s*:\s*(.+?)\]', resposta)
            if m_tarefa:
                try:
                    adicionar_tarefa(m_tarefa.group(1).strip())
                    resposta = re.sub(r'\[TAREFA[^\]]*\]', '', resposta).strip()
                except Exception:
                    pass

            m_concluir = re.search(r'\[CONCLUIR\s*:\s*(.+?)\]', resposta)
            if m_concluir:
                try:
                    concluir_tarefa_por_texto(m_concluir.group(1).strip())
                    resposta = re.sub(r'\[CONCLUIR[^\]]*\]', '', resposta).strip()
                except Exception:
                    pass

            m_memoria = re.search(r'\[MEMORIA\s*:\s*(.+?)\]', resposta)
            if m_memoria:
                try:
                    adicionar_memoria(m_memoria.group(1).strip())
                    resposta = re.sub(r'\[MEMORIA[^\]]*\]', '', resposta).strip()
                except Exception:
                    pass

            if not resposta.strip():
                resposta = "👍"

            bot.reply_to(msg, resposta)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    global bot

    print("=" * 54)
    print("   🌟 ORION — Assistente Pessoal (Telegram)")
    print("=" * 54)

    carregar_config()

    token = config["telegram"].get("token", "")
    if "SEU_TOKEN" in token or not token:
        print("\n❌ TELEGRAM_TOKEN não configurado!")
        sys.exit(1)

    api_key = config.get("groq", {}).get("api_key", "")
    if not api_key or "SUA_KEY" in api_key:
        print("\n⚠️  GROQ_API_KEY não configurada.")

    if _USA_SB:
        print("\n✅ Persistência: Supabase")
    else:
        print("\n⚠️  Persistência: arquivos locais (dados perdidos ao reiniciar)")

    bot = telebot.TeleBot(token)
    registrar_handlers()
    configurar_agenda()

    agendador = threading.Thread(target=thread_agenda, daemon=True)
    agendador.start()

    chat_id = config["telegram"].get("chat_id", "")
    if not chat_id:
        print("\n⚠️  Envie /start ao bot para vincular seu chat.")
    else:
        print(f"\n✅ Chat vinculado: {chat_id}")

    print("\n✅ Orion rodando!\n")

    bot.infinity_polling(timeout=30, long_polling_timeout=20)


if __name__ == "__main__":
    main()
